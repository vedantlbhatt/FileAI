import os
from typing import List, Tuple
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from striprtf.striprtf import rtf_to_text
import subprocess
from PIL import Image
import torch
from transformers import CLIPProcessor, CLIPModel

# Load YOLOv5 model with torch hub
try:
    yolo_model = torch.hub.load('ultralytics/yolov5', 'yolov5s', pretrained=True)
except Exception as e:
    print(f"Error loading YOLOv5 model: {e}")
    yolo_model = None

def parse_rtf_unrtf(file_path):
    result = subprocess.run(['unrtf', '--text', file_path], capture_output=True, text=True)
    if result.returncode == 0:
        return result.stdout.strip()
    else:
        print(f"UnRTF failed: {result.stderr}")
        return ""

def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return '\n'.join([line for line in lines if line])

def parse_txt(filepath: str) -> str:
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        print(f"Skipping non-UTF8 or binary file: {filepath}")
        return ""

def parse_pdf(filepath: str) -> str:
    import PyPDF2
    text = ""
    with open(filepath, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def parse_docx(filepath: str) -> str:
    import docx
    doc = docx.Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def parse_file(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.txt':
        raw_text = parse_txt(filepath)
    elif ext == '.pdf':
        raw_text = parse_pdf(filepath)
    elif ext == '.docx':
        raw_text = parse_docx(filepath)
    elif ext == '.rtf':
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                raw_rtf = f.read()
            raw_text = parse_rtf_unrtf(filepath)
        except Exception as e:
            print(f"Error parsing RTF file {filepath}: {e}")
            raw_text = ""
    else:
        raw_text = ""
    return clean_text(raw_text)

def get_all_files_and_parse(directory: str,
                            text_exts=('.txt', '.pdf', '.docx', '.rtf'),
                            image_exts=('.jpg', '.jpeg', '.png', '.bmp', '.gif')) -> Tuple[List[str], List[str], List[str]]:
    text_texts = []
    text_paths = []
    image_paths = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            path = os.path.join(root, file)
            if file.lower().endswith(text_exts):
                try:
                    text = parse_file(path)
                    if text:
                        text_texts.append(text)
                        text_paths.append(path)
                except Exception as e:
                    print(f"Error parsing {path}: {e}")
            elif file.lower().endswith(image_exts):
                image_paths.append(path)
    return text_texts, text_paths, image_paths

def embed_and_index_texts(texts: List[str]):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    embeddings = model.encode(texts, convert_to_tensor=False)
    embeddings_np = np.array(embeddings).astype('float32')
    dimension = embeddings_np.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings_np)
    return index, model, embeddings_np.shape[0]

clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")

def embed_text_clip(text: str) -> np.ndarray:
    inputs = clip_processor(text=[text], return_tensors="pt", padding=True)
    with torch.no_grad():
        text_features = clip_model.get_text_features(**inputs)
    text_features /= text_features.norm(p=2, dim=-1, keepdim=True)
    return text_features[0].cpu().numpy()

def detect_and_crop_objects(image_path: str) -> List[Image.Image]:
    img = Image.open(image_path).convert('RGB')
    if yolo_model is None:
        return [img]
    results = yolo_model(img)
    crops = []
    for *box, conf, cls in results.xyxy[0]:
        x1, y1, x2, y2 = map(int, box)
        crop = img.crop((x1, y1, x2, y2))
        crops.append(crop)
    return crops if crops else [img]

def embed_images_with_object_detection(image_paths: List[str]) -> Tuple[np.ndarray, List[Tuple[int,int]]]:
    all_embeddings = []
    image_to_region = []
    for img_idx, img_path in enumerate(image_paths):
        try:
            crops = detect_and_crop_objects(img_path)
            for region_idx, crop_img in enumerate(crops):
                inputs = clip_processor(images=crop_img, return_tensors="pt")
                with torch.no_grad():
                    emb = clip_model.get_image_features(**inputs)
                emb /= emb.norm(p=2, dim=-1, keepdim=True)
                emb_np = emb[0].cpu().numpy()
                all_embeddings.append(emb_np)
                image_to_region.append((img_idx, region_idx))
        except Exception as e:
            print(f"Error processing image {img_path}: {e}")
    if not all_embeddings:
        print("No image embeddings created.")
        return np.array([]), []
    embeddings_np = np.array(all_embeddings).astype('float32')
    return embeddings_np, image_to_region

def build_image_index(embeddings_np: np.ndarray):
    dimension = embeddings_np.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings_np)
    return index

def semantic_search_images(index, image_paths, image_to_region, query: str, k: int=5):
    text_emb = embed_text_clip(query).astype('float32')
    distances, indices = index.search(text_emb.reshape(1, -1), k)
    results = []
    seen_images = set()
    for dist, idx in zip(distances[0], indices[0]):
        img_idx, region_idx = image_to_region[idx]
        file_path = image_paths[img_idx]
        score = 1 / (1 + dist)
        if file_path not in seen_images:
            results.append({'file_path': file_path, 'semantic_score': score, 'distance': dist, 'region_index': region_idx})
            seen_images.add(file_path)
    return results

def keyword_score(text: str, query: str) -> float:
    query_words = set(query.lower().split())
    text_words = set(text.lower().split())
    common = query_words.intersection(text_words)
    return len(common) / len(query_words) if query_words else 0

def hybrid_semantic_search(index, model, texts, file_paths, query: str, k: int=5, alpha: float = 0.5):
    k = min(k, len(texts))
    query_embedding = model.encode([query], convert_to_tensor=False)
    query_embedding_np = np.array(query_embedding).astype('float32')
    distances, indices = index.search(query_embedding_np, k)
    results = []
    max_distance = max(distances[0]) if distances[0].size > 0 else 1e-6
    for dist, idx in zip(distances[0], indices[0]):
        sem_score = 1 - dist / (max_distance + 1e-12)
        kw_score = keyword_score(texts[idx], query)
        combined_score = alpha * sem_score + (1 - alpha) * kw_score
        results.append({'file_path': file_paths[idx], 'semantic_score': sem_score, 'keyword_score': kw_score, 'combined_score': combined_score})
    results = sorted(results, key=lambda x: x['combined_score'], reverse=True)
    return results[:k]

if __name__ == "__main__":
    directory_path = '/Users/vedantbhatt/source'
    query_text = "family"
    print("Parsing text files...")
    all_texts, all_file_paths, image_paths = get_all_files_and_parse(directory_path)
from typing import List, Tuple
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from striprtf.striprtf import rtf_to_text
import subprocess
from PIL import Image
import torch
from transformers import CLIPProcessor, CLIPModel
from transformers import BlipProcessor, BlipForConditionalGeneration


# Load YOLOv5 model with torch hub
try:
    yolo_model = torch.hub.load('ultralytics/yolov5', 'yolov5s', pretrained=True)
except Exception as e:
    print(f"Error loading YOLOv5 model: {e}")
    yolo_model = None


def parse_rtf_unrtf(file_path):
    result = subprocess.run(['unrtf', '--text', file_path], capture_output=True, text=True)
    if result.returncode == 0:
        return result.stdout.strip()
    else:
        print(f"UnRTF failed: {result.stderr}")
        return ""


def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return '\n'.join([line for line in lines if line])


def parse_txt(filepath: str) -> str:
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        print(f"Skipping non-UTF8 or binary file: {filepath}")
        return ""


def parse_pdf(filepath: str) -> str:
    import PyPDF2
    text = ""
    with open(filepath, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def parse_docx(filepath: str) -> str:
    import docx
    doc = docx.Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)


def parse_file(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.txt':
        raw_text = parse_txt(filepath)
    elif ext == '.pdf':
        raw_text = parse_pdf(filepath)
    elif ext == '.docx':
        raw_text = parse_docx(filepath)
    elif ext == '.rtf':
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                raw_rtf = f.read()
            raw_text = parse_rtf_unrtf(filepath)
        except Exception as e:
            print(f"Error parsing RTF file {filepath}: {e}")
            raw_text = ""
    else:
        raw_text = ""
    return clean_text(raw_text)


def get_all_files_and_parse(directory: str,
                            text_exts=('.txt', '.pdf', '.docx', '.rtf'),
                            image_exts=('.jpg', '.jpeg', '.png', '.bmp', '.gif')) -> Tuple[List[str], List[str], List[str]]:
    text_texts = []
    text_paths = []
    image_paths = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            path = os.path.join(root, file)
            if file.lower().endswith(text_exts):
                try:
                    text = parse_file(path)
                    if text:
                        text_texts.append(text)
                        text_paths.append(path)
                except Exception as e:
                    print(f"Error parsing {path}: {e}")
            elif file.lower().endswith(image_exts):
                image_paths.append(path)
    return text_texts, text_paths, image_paths


def embed_and_index_texts(texts: List[str]):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    embeddings = model.encode(texts, convert_to_tensor=False)
    embeddings_np = np.array(embeddings).astype('float32')
    dimension = embeddings_np.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings_np)
    return index, model, embeddings_np.shape[0]


clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")

blip_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
blip_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")


def embed_text_clip(text: str) -> np.ndarray:
    inputs = clip_processor(text=[text], return_tensors="pt", padding=True)
    with torch.no_grad():
        text_features = clip_model.get_text_features(**inputs)
    text_features /= text_features.norm(p=2, dim=-1, keepdim=True)
    return text_features[0].cpu().numpy()


def generate_caption(image: Image.Image) -> str:
    inputs = blip_processor(image, return_tensors="pt")
    out = blip_model.generate(**inputs)
    caption = blip_processor.decode(out[0], skip_special_tokens=True)
    return caption


def detect_and_crop_objects(image_path: str) -> List[Image.Image]:
    img = Image.open(image_path).convert('RGB')
    if yolo_model is None:
        return [img]
    results = yolo_model(img)
    crops = []
    for *box, conf, cls in results.xyxy[0]:
        x1, y1, x2, y2 = map(int, box)
        crop = img.crop((x1, y1, x2, y2))
        crops.append(crop)
    return crops if crops else [img]


def embed_images_with_object_detection(image_paths: List[str]) -> Tuple[np.ndarray, List[Tuple[int,int]], List[str]]:
    all_embeddings = []
    image_to_region = []
    captions = []
    for img_idx, img_path in enumerate(image_paths):
        try:
            crops = detect_and_crop_objects(img_path)
            for region_idx, crop_img in enumerate(crops):
                inputs = clip_processor(images=crop_img, return_tensors="pt")
                with torch.no_grad():
                    emb = clip_model.get_image_features(**inputs)
                emb /= emb.norm(p=2, dim=-1, keepdim=True)
                emb_np = emb[0].cpu().numpy()
                all_embeddings.append(emb_np)
                image_to_region.append((img_idx, region_idx))
                # Generate captions for reranking
                captions.append(generate_caption(crop_img))
        except Exception as e:
            print(f"Error processing image {img_path}: {e}")
            captions.append("")  # fallback
    if not all_embeddings:
        print("No image embeddings created.")
        return np.array([]), [], []
    embeddings_np = np.array(all_embeddings).astype('float32')
    return embeddings_np, image_to_region, captions


def build_image_index(embeddings_np: np.ndarray):
    dimension = embeddings_np.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings_np)
    return index


def semantic_search_images(index, image_paths, image_to_region, captions, query: str, k: int=5):
    text_emb = embed_text_clip(query).astype('float32')
    distances, indices = index.search(text_emb.reshape(1, -1), k)
    results = []
    for dist, idx in zip(distances[0], indices[0]):
        img_idx, region_idx = image_to_region[idx]
        file_path = image_paths[img_idx]
        caption = captions[idx]
        score = 1 / (1 + dist)
        results.append({'file_path': file_path,
                        'semantic_score': score,
                        'distance': dist,
                        'region_index': region_idx,
                        'caption': caption})
    # Re-rank results by caption similarity to query
    reranked_results = rerank_by_caption_similarity(results, query)
    return reranked_results


def rerank_by_caption_similarity(image_results, query: str):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    caption_texts = [r['caption'] for r in image_results]
    caption_embeddings = model.encode(caption_texts)
    query_embedding = model.encode([query])[0]

    def cosine_sim(a, b):
        return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-10)

    scores = [cosine_sim(query_embedding, cap_emb) for cap_emb in caption_embeddings]
    for res, score in zip(image_results, scores):
        res['rerank_score'] = score
    reranked = sorted(image_results, key=lambda x: x['rerank_score'], reverse=True)
    return reranked


def keyword_score(text: str, query: str) -> float:
    query_words = set(query.lower().split())
    text_words = set(text.lower().split())
    common = query_words.intersection(text_words)
    return len(common) / len(query_words) if query_words else 0


def hybrid_semantic_search(index, model, texts, file_paths, query: str, k: int=5, alpha: float = 0.5):
    k = min(k, len(texts))
    query_embedding = model.encode([query], convert_to_tensor=False)
    query_embedding_np = np.array(query_embedding).astype('float32')
    distances, indices = index.search(query_embedding_np, k)
    results = []
    max_distance = max(distances[0]) if distances[0].size > 0 else 1e-6
    for dist, idx in zip(distances[0], indices[0]):
        sem_score = 1 - dist / (max_distance + 1e-12)
        kw_score = keyword_score(texts[idx], query)
        combined_score = alpha * sem_score + (1 - alpha) * kw_score
        results.append({'file_path': file_paths[idx], 'semantic_score': sem_score, 'keyword_score': kw_score, 'combined_score': combined_score})
    results = sorted(results, key=lambda x: x['combined_score'], reverse=True)
    return results[:k]


if __name__ == "__main__":
    directory_path = '/Users/vedantbhatt/source'
    query_text = "family"

    print("Parsing text files...")
    all_texts, all_file_paths, image_paths = get_all_files_and_parse(directory_path)
    print(f"Parsed {len(all_texts)} text files and found {len(image_paths)} images.")

    print("Embedding texts and building text index...")
    text_index, text_model, _ = embed_and_index_texts(all_texts)

    print("Detecting objects, embedding cropped image regions and building image index...")
    image_embeddings_np, image_to_region, captions = embed_images_with_object_detection(image_paths)
    if image_embeddings_np.size == 0:
        print("No images to index.")
    else:
        image_index = build_image_index(image_embeddings_np)
        print(f"Searching texts for query: {query_text}")
        text_results = hybrid_semantic_search(text_index, text_model, all_texts, all_file_paths, query_text, k=5, alpha=0.7)

        print(f"Searching images for query: {query_text}")
        image_results = semantic_search_images(image_index, image_paths, image_to_region, captions, query_text, k=20)

        print("\nText Search Results:")
        for res in text_results:
            print(f"File: {res['file_path']}")
            print(f"Combined Score: {res['combined_score']:.4f} (Semantic: {res['semantic_score']:.4f}, Keyword: {res['keyword_score']:.4f})")
            print()
        print("\nImage Search Results:")
        for res in image_results:
            print(f"File: {res['file_path']} (Region: {res['region_index']})")
            print(f"Semantic Score: {res['semantic_score']:.4f}, Distance: {res['distance']:.4f}, Rerank Score: {res['rerank_score']:.4f}")
            print(f"Caption: {res['caption']}")
            print()

    print(f"Parsed {len(all_texts)} text files and found {len(image_paths)} images.")
    print("Embedding texts and building text index...")
    text_index, text_model, _ = embed_and_index_texts(all_texts)
    print("Detecting objects, embedding cropped image regions and building image index...")
    image_embeddings_np, image_to_region = embed_images_with_object_detection(image_paths)
    if image_embeddings_np.size == 0:
        print("No images to index.")
    else:
        image_index = build_image_index(image_embeddings_np)
        print(f"Searching texts for query: {query_text}")
        text_results = hybrid_semantic_search(text_index, text_model, all_texts, all_file_paths, query_text, k=5, alpha=0.7)
        print(f"Searching images for query: {query_text}")
        image_results = semantic_search_images(image_index, image_paths, image_to_region, query_text, k=20)
        for res in image_results:
            print(f"File: {res['file_path']} (Region: {res['region_index']})")
            print(f"Semantic Score: {res['semantic_score']:.4f}, Distance: {res['distance']:.4f}, Rerank Score: {res['rerank_score']:.4f}")
            print(f"Caption: {res['caption']}")
            print()


        print("\nText Search Results:")
        for res in text_results:
            print(f"File: {res['file_path']}")
            print(f"Combined Score: {res['combined_score']:.4f} (Semantic: {res['semantic_score']:.4f}, Keyword: {res['keyword_score']:.4f})")
            print()
        print("\nImage Search Results:")
        for res in image_results:
            print(f"File: {res['file_path']} (Region: {res['region_index']})")
            print(f"Semantic Score: {res['semantic_score']:.4f}, Distance: {res['distance']:.4f}")
            print()
